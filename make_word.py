#coding:utf-8
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE

import os
import sys

def get_docx_file(report_items=None):
	try:
		doc_name = u"{}-{}-{}-个人绩效日报.docx".format(report_items["department"], report_items["name"], report_items["date"].replace("-",""))
		#doc_name = doc_name.encode("gbk")
		#file_path = os.path.join(os.getcwd(),"Daily_Report")
		file_path = os.path.join(os.path.split(sys.argv[0])[0],"Daily_Report")

		if not os.path.exists(file_path):
			os.mkdir(file_path)
		daily_report_path = os.path.join(file_path,doc_name)

		f = open(daily_report_path, 'w')
		f.close()
		return daily_report_path
	except Exception as e:
		print e
		return False

def built_docx(report_items=None):
    daily_report_path = get_docx_file(report_items)

    try:
        document = Document(docx="default.docx")

        head_style = document.styles.add_style("DailyReport",style_type=WD_STYLE_TYPE.PARAGRAPH)
        # style = document.styles['Normal']
        head_style.font.name = 'Arial Unicode MS'
        head_style.font.size = Pt(24)

        table_style = document.styles.add_style("DailyReportTable",style_type=WD_STYLE_TYPE.TABLE)
        table_style.font.name = u'宋体'
        table_style.font.size = Pt(10)
    
        p = document.add_paragraph(u'大数据和平台总体组个人绩效日报',style="DailyReport")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=5, cols=6, style="DailyReportTable")
#        table.style.quick_style = True
#        table.style.hidden = False
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
	
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'部门'
        hdr_cells[2].text = u'姓名'
        hdr_cells[4].text = u'日期'

        table.cell(1,0).text = u'''
        今
        日
        任
        务
        完
        成
        情
        况
        '''

        table.cell(2, 0).text = u'''
        未
        完
        成
        或
        遇
        到
        问
        题
        '''

        table.cell(3, 0).text = u'''
        明
        日
        计
        划
        '''

        table.cell(4, 0).text = u'''
        收
        获
        感
        悟
        '''
        for i in range(1,5):
            for j in range(1,5)[::-1]:
                table.rows[i].cells[j].merge(table.rows[i].cells[j+1])

        table.cell(0,1).text = report_items["department"]
        table.cell(0,3).text = report_items["name"]
        table.cell(0,5).text = report_items["date"]
        table.cell(1,1).text = report_items["today_data"]
        table.cell(2,1).text = report_items["unfinished_data"]
        table.cell(3,1).text = report_items["tomorrow_data"]
        table.cell(4,1).text = report_items["think_data"]

        document.save(daily_report_path)
        return daily_report_path
    except Exception as e:
        print "6",e
        #return False


if __name__=="__main__":
    report_items = {
    "department":u"大数据实验室",
    "name":u"谭勇",
    "date":"2016-06-16",
    "today_data":u"""
    1.云南公安URL数据分析
    今天是在做数据提取
    2.协助铁路公安FEA在公司部署测试，
    """,
    "unfinished_data":u"""
    云南使用的是esql3，没有导出功能。不能将数据到出到本地分析。
    """,
    "tomorrow_data":u"""
    云南数据分析，先连到云南的ESQL上看数据结构，然后找孙影帮忙看下怎样可以导出数据
    """,
    "think_data":u""
    }

    built_docx(report_items)